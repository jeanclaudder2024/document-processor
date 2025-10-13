"""
Create a sample Word template for testing
"""

from docx import Document
from docx.shared import Inches

def create_sample_template():
    """Create a sample vessel document template"""
    
    # Create a new document
    doc = Document()
    
    # Add title
    title = doc.add_heading('Vessel Document', 0)
    
    # Add vessel information section
    doc.add_heading('Vessel Information', level=1)
    
    # Create a table for vessel details
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    
    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Field'
    hdr_cells[1].text = 'Value'
    
    # Add vessel data rows
    vessel_data = [
        ('Vessel Name', '{{vessel_name}}'),
        ('IMO Number', '{{imo}}'),
        ('Vessel Type', '{{vessel_type}}'),
        ('Flag', '{{flag}}'),
        ('MMSI', '{{mmsi}}'),
        ('Call Sign', '{{callsign}}'),
        ('Built Year', '{{built}}'),
        ('Deadweight', '{{deadweight}}'),
        ('Length', '{{length}}'),
        ('Width', '{{width}}'),
        ('Draught', '{{draught}}'),
        ('Gross Tonnage', '{{gross_tonnage}}'),
        ('Engine Power', '{{engine_power}}'),
        ('Crew Size', '{{crew_size}}'),
    ]
    
    for field, placeholder in vessel_data:
        row_cells = table.add_row().cells
        row_cells[0].text = field
        row_cells[1].text = placeholder
    
    # Add commercial information section
    doc.add_heading('Commercial Information', level=1)
    
    # Create another table for commercial details
    comm_table = doc.add_table(rows=1, cols=2)
    comm_table.style = 'Table Grid'
    
    # Header row
    hdr_cells = comm_table.rows[0].cells
    hdr_cells[0].text = 'Field'
    hdr_cells[1].text = 'Value'
    
    # Add commercial data rows
    commercial_data = [
        ('Owner', '{{owner_name}}'),
        ('Operator', '{{operator_name}}'),
        ('Buyer', '{{buyer_name}}'),
        ('Seller', '{{seller_name}}'),
        ('Cargo Type', '{{cargo_type}}'),
        ('Cargo Quantity', '{{cargo_quantity}}'),
        ('Oil Type', '{{oil_type}}'),
        ('Oil Source', '{{oil_source}}'),
        ('Departure Port', '{{departure_port}}'),
        ('Destination Port', '{{destination_port}}'),
        ('Loading Port', '{{loading_port}}'),
        ('Departure Date', '{{departure_date}}'),
        ('Arrival Date', '{{arrival_date}}'),
        ('ETA', '{{eta}}'),
        ('Current Region', '{{current_region}}'),
        ('Status', '{{status}}'),
        ('Speed', '{{speed}}'),
        ('Course', '{{course}}'),
        ('Deal Value', '{{deal_value}}'),
        ('Price', '{{price}}'),
        ('Market Price', '{{market_price}}'),
    ]
    
    for field, placeholder in commercial_data:
        row_cells = comm_table.add_row().cells
        row_cells[0].text = field
        row_cells[1].text = placeholder
    
    # Add a paragraph with some text
    doc.add_paragraph()
    doc.add_paragraph('This document contains information about the vessel {{vessel_name}} with IMO number {{imo}}. The vessel is currently {{status}} in the {{current_region}} region.')
    
    # Add another paragraph
    doc.add_paragraph('The vessel is carrying {{cargo_quantity}} barrels of {{cargo_type}} from {{departure_port}} to {{destination_port}}. The estimated arrival time is {{eta}}.')
    
    # Save the document
    doc.save('templates/Sample_Vessel_Document.docx')
    print("Sample template created: templates/Sample_Vessel_Document.docx")

if __name__ == "__main__":
    create_sample_template()

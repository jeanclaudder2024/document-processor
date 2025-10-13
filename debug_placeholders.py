from docx import Document
import os
import re

def debug_template_placeholders(filename):
    print(f'\n=== DEBUGGING {filename} ===')
    file_path = os.path.join('./templates', filename)
    doc = Document(file_path)
    
    # Extract all text content
    full_text = ''
    for paragraph in doc.paragraphs:
        full_text += paragraph.text + '\n'
    
    # Also check tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    full_text += paragraph.text + '\n'
    
    # Find all possible placeholder patterns
    patterns = [
        r'\{\{([^}]+)\}\}',  # {{placeholder}}
        r'\{([^}]+)\}',      # {placeholder}
        r'\[([^\]]+)\]',     # [placeholder]
        r'\[\[([^\]]+)\]\]', # [[placeholder]]
        r'%([^%]+)%',        # %placeholder%
        r'<([^>]+)>',        # <placeholder>
        r'__([^_]+)__',      # __placeholder__
        r'##([^#]+)##',      # ##placeholder##
    ]
    
    print(f'Full text length: {len(full_text)} characters')
    print(f'First 500 characters: {full_text[:500]}')
    
    all_placeholders = []
    for i, pattern in enumerate(patterns):
        matches = re.findall(pattern, full_text)
        if matches:
            print(f'Pattern {i+1} ({pattern}): Found {len(matches)} matches')
            for match in matches[:5]:  # Show first 5 matches
                print(f'  - "{match}"')
        all_placeholders.extend(matches)
    
    # Show unique placeholders
    unique_placeholders = list(set(all_placeholders))
    print(f'\nTotal unique placeholders found: {len(unique_placeholders)}')
    for p in sorted(unique_placeholders):
        print(f'  - "{p}"')

# Test with one template
debug_template_placeholders('Sample_Vessel_Document.docx')

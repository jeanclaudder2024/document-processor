from main import find_placeholders
from docx import Document
import os

def test_placeholder_finding():
    print("Testing placeholder finding...")
    
    # Test with Sample_Vessel_Document.docx
    file_path = os.path.join('./templates', 'Sample_Vessel_Document.docx')
    doc = Document(file_path)
    
    # Extract all text content
    full_text = ''
    for paragraph in doc.paragraphs:
        full_text += paragraph.text + '\n'
    
    # Also check tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    full_text += paragraph.text + '\n'
    
    placeholders = find_placeholders(full_text)
    print(f"Found {len(placeholders)} placeholders:")
    for p in placeholders:
        print(f"  - {p}")

if __name__ == "__main__":
    test_placeholder_finding()

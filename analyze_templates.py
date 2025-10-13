from docx import Document
import os
import re

def analyze_template(filename):
    print(f'\n=== ANALYZING {filename} ===')
    file_path = os.path.join('./templates', filename)
    doc = Document(file_path)
    
    # Extract all text content
    full_text = ''
    for paragraph in doc.paragraphs:
        full_text += paragraph.text + '\n'
    
    # Also check tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    full_text += paragraph.text + '\n'
    
    # Find all possible placeholder patterns
    patterns = [
        r'\{\{([^}]+)\}\}',  # {{placeholder}}
        r'\{([^}]+)\}',      # {placeholder}
        r'\[([^\]]+)\]',     # [placeholder]
        r'\[\[([^\]]+)\]\]', # [[placeholder]]
        r'%([^%]+)%',        # %placeholder%
        r'<([^>]+)>',        # <placeholder>
        r'__([^_]+)__',      # __placeholder__
        r'##([^#]+)##',      # ##placeholder##
    ]
    
    all_placeholders = []
    for pattern in patterns:
        matches = re.findall(pattern, full_text)
        all_placeholders.extend(matches)
    
    # Show unique placeholders
    unique_placeholders = list(set(all_placeholders))
    print(f'Found {len(unique_placeholders)} unique placeholders:')
    for p in sorted(unique_placeholders):
        print(f'  - "{p}"')
    
    # Show some context around placeholders
    print('\nContext examples:')
    lines = full_text.split('\n')
    for line in lines:
        if any(pattern in line for pattern in ['{{', '{', '[', '%', '<', '__', '##']):
            if line.strip():
                print(f'  "{line.strip()}"')

# Analyze all templates
templates = [f for f in os.listdir('./templates') if f.endswith('.docx')]
for template in templates:
    analyze_template(template)

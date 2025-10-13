from docx import Document
import os
import re

def analyze_template(filename):
    print(f'\n=== ANALYZING {filename} ===')
    file_path = os.path.join('./templates', filename)
    doc = Document(file_path)
    
    # Extract all text content
    full_text = ''
    for paragraph in doc.paragraphs:
        full_text += paragraph.text + '\n'
    
    # Also check tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    full_text += paragraph.text + '\n'
    
    # Find all possible placeholder patterns
    patterns = [
        r'\{\{([^}]+)\}\}',  # {{placeholder}}
        r'\{([^}]+)\}',      # {placeholder}
        r'\[([^\]]+)\]',     # [placeholder]
        r'\[\[([^\]]+)\]\]', # [[placeholder]]
        r'%([^%]+)%',        # %placeholder%
        r'<([^>]+)>',        # <placeholder>
        r'__([^_]+)__',      # __placeholder__
        r'##([^#]+)##',      # ##placeholder##
    ]
    
    all_placeholders = []
    for pattern in patterns:
        matches = re.findall(pattern, full_text)
        all_placeholders.extend(matches)
    
    # Clean and filter placeholders
    clean_placeholders = []
    for p in all_placeholders:
        p_clean = p.strip()
        if p_clean and len(p_clean) < 100:  # Filter out very long strings
            clean_placeholders.append(p_clean)
    
    # Show unique placeholders
    unique_placeholders = list(set(clean_placeholders))
    print(f'Found {len(unique_placeholders)} unique placeholders:')
    for p in sorted(unique_placeholders):
        print(f'  - {p}')

# Analyze all templates
templates = [f for f in os.listdir('./templates') if f.endswith('.docx')]
for template in templates:
    analyze_template(template)

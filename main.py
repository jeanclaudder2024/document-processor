import os
from pathlib import Path

# Load .env from document-processor and project root (for local dev)
def _load_dotenv():
    try:
        from dotenv import load_dotenv
        for path in [Path(__file__).resolve().parent / ".env", Path(__file__).resolve().parent.parent / ".env"]:
            if path.exists():
                try:
                    load_dotenv(path)
                except (UnicodeDecodeError, Exception):
                    pass
        try:
            load_dotenv()
        except (UnicodeDecodeError, Exception):
            pass
    except ImportError:
        pass
    except Exception:
        pass

_load_dotenv()

from fastapi import FastAPI, UploadFile, File, Form, HTTPException, Query
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, Response, StreamingResponse
from docx import Document
from docx.shared import Pt
from supabase import create_client
import re
import json
import base64
from io import BytesIO
from typing import Optional, List, Dict, Any
import uuid
from datetime import datetime, timedelta
import random
from pydantic import BaseModel
import aiofiles
import pandas as pd
from openai import OpenAI
import subprocess
import tempfile
from pdf2image import convert_from_path
import img2pdf
from PIL import Image

try:
    from docx2pdf import convert as docx2pdf_convert
    DOCX2PDF_AVAILABLE = True
except ImportError:
    DOCX2PDF_AVAILABLE = False
    print("[WARNING] docx2pdf not available, will use LibreOffice directly")

app = FastAPI(title="PetroDealHub Document Processor", version="1.0.0")

from starlette.middleware.base import BaseHTTPMiddleware
from starlette.requests import Request
from starlette.responses import Response

class OpenCORSMiddleware(BaseHTTPMiddleware):
    """Allow any origin to connect - supports Lovable Settings tab custom domains."""
    async def dispatch(self, request: Request, call_next):
        origin = request.headers.get("origin", "")
        
        if request.method == "OPTIONS":
            response = Response(status_code=200)
            if origin:
                response.headers["Access-Control-Allow-Origin"] = origin
                response.headers["Access-Control-Allow-Credentials"] = "true"
            else:
                response.headers["Access-Control-Allow-Origin"] = "*"
            response.headers["Access-Control-Allow-Methods"] = "GET, POST, PUT, DELETE, OPTIONS, PATCH"
            response.headers["Access-Control-Allow-Headers"] = request.headers.get(
                "Access-Control-Request-Headers", "Content-Type, Authorization, X-Requested-With"
            )
            response.headers["Access-Control-Max-Age"] = "86400"
            response.headers["Vary"] = "Origin"
            return response
        
        response = await call_next(request)
        if origin:
            response.headers["Access-Control-Allow-Origin"] = origin
            response.headers["Access-Control-Allow-Credentials"] = "true"
        else:
            response.headers["Access-Control-Allow-Origin"] = "*"
        response.headers["Vary"] = "Origin"
        
        return response

app.add_middleware(OpenCORSMiddleware)

SUPABASE_URL = os.getenv("SUPABASE_URL")
SUPABASE_SERVICE_ROLE_KEY = os.getenv("SUPABASE_SERVICE_ROLE_KEY")
SUPABASE_KEY = os.getenv("SUPABASE_KEY")
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")

supabase = None
openai_client = None

supabase_key_to_use = SUPABASE_SERVICE_ROLE_KEY or SUPABASE_KEY
if SUPABASE_URL and supabase_key_to_use:
    supabase = create_client(SUPABASE_URL, supabase_key_to_use)

if OPENAI_API_KEY:
    openai_client = OpenAI(api_key=OPENAI_API_KEY)

TEMPLATES_DIR = "templates"
CSV_DIR = "csv_data"
SETTINGS_DIR = "settings"
FONTS_DIR = "fonts"

os.makedirs(TEMPLATES_DIR, exist_ok=True)
os.makedirs(CSV_DIR, exist_ok=True)
os.makedirs(SETTINGS_DIR, exist_ok=True)

fonts_home = os.path.expanduser("~/.fonts")
os.makedirs(fonts_home, exist_ok=True)
if os.path.isdir(FONTS_DIR):
    import shutil
    for f in os.listdir(FONTS_DIR):
        if f.endswith(".ttf"):
            src = os.path.join(FONTS_DIR, f)
            dst = os.path.join(fonts_home, f)
            if not os.path.exists(dst):
                shutil.copy2(src, dst)
    
    fc_dir = os.path.expanduser("~/.config/fontconfig/conf.d")
    os.makedirs(fc_dir, exist_ok=True)
    alias_conf = os.path.join(fc_dir, "30-font-aliases.conf")
    with open(alias_conf, "w") as f:
        f.write("""<?xml version="1.0"?>
<!DOCTYPE fontconfig SYSTEM "fonts.dtd">
<fontconfig>
  <alias><family>Arial</family><prefer><family>Liberation Sans</family></prefer></alias>
  <alias><family>Courier New</family><prefer><family>Liberation Mono</family></prefer></alias>
  <alias><family>Times New Roman</family><prefer><family>Liberation Serif</family></prefer></alias>
  <alias><family>Calibri</family><prefer><family>Carlito</family></prefer></alias>
  <alias><family>Helvetica</family><prefer><family>Liberation Sans</family></prefer></alias>
  <alias><family>OCR-B 10 BT</family><prefer><family>OCR-B-10-BT</family></prefer></alias>
  <alias><family>OCR B</family><prefer><family>OCRB</family></prefer></alias>
</fontconfig>""")
    
    os.popen("fc-cache -f ~/.fonts/ 2>/dev/null")
    ttf_count = len([f for f in os.listdir(FONTS_DIR) if f.endswith(".ttf")])
    print(f"[FONTS] Installed {ttf_count} font files from {FONTS_DIR}/ to ~/.fonts/")
    print("[FONTS] Font aliases configured: Arial->Liberation Sans, Courier New->Liberation Mono, OCR-B 10 BT->OCR-B-10-BT")

PLACEHOLDER_PATTERNS = [
    r'\{\{([^}]+)\}\}',
    r'\{([^}]+)\}',
    r'\[\[([^\]]+)\]\]',
    r'\[([^\]]+)\]',
    r'%([^%]+)%',
    r'<<([^>]+)>>',
    r'##([^#]+)##',
]

DATABASE_TABLES = [
    "vessels", "ports", "refineries", "buyer_companies", "seller_companies",
    "oil_products", "broker_profiles", "companies", "buyer_company_bank_accounts",
    "seller_company_bank_accounts"
]

PREFIX_TO_TABLE = {
    "vessel_": "vessels",
    "port_": "ports",
    "departure_port_": "ports",
    "destination_port_": "ports",
    "buyer_": "buyer_companies",
    "seller_": "seller_companies",
    "refinery_": "refineries",
    "product_": "oil_products",
    "buyer_bank_": "buyer_company_bank_accounts",
    "seller_bank_": "seller_company_bank_accounts",
    "company_": "companies",
}


class TemplateMetadata(BaseModel):
    display_name: Optional[str] = None
    description: Optional[str] = None
    font_family: Optional[str] = None
    font_size: Optional[int] = None
    plan_ids: Optional[List[str]] = None


class PlaceholderSettingsRequest(BaseModel):
    template_name: str
    template_id: Optional[str] = None
    settings: Dict[str, Any]


class ProcessDocumentRequest(BaseModel):
    template_name: str
    vessel_id: Optional[int] = None
    buyer_id: Optional[str] = None
    seller_id: Optional[str] = None
    product_id: Optional[str] = None
    refinery_id: Optional[str] = None
    departure_port_id: Optional[int] = None
    destination_port_id: Optional[int] = None
    buyer_bank_id: Optional[str] = None
    seller_bank_id: Optional[str] = None
    output_format: str = "base64"


class GenerateDocumentRequest(BaseModel):
    template_name: str
    vessel_imo: Optional[str] = None
    vessel_id: Optional[int] = None
    buyer_id: Optional[str] = None
    seller_id: Optional[str] = None
    product_id: Optional[str] = None
    refinery_id: Optional[str] = None
    departure_port_id: Optional[int] = None
    destination_port_id: Optional[int] = None
    buyer_bank_id: Optional[str] = None
    seller_bank_id: Optional[str] = None


class PlanUpdateRequest(BaseModel):
    plan_id: str
    plan_data: Dict[str, Any]


def is_valid_uuid(val: str) -> bool:
    """Check if a string is a valid UUID format"""
    if not val:
        return False
    try:
        uuid.UUID(str(val))
        return True
    except (ValueError, AttributeError):
        return False


def extract_placeholders_from_text(text: str) -> List[str]:
    placeholders = []
    for pattern in PLACEHOLDER_PATTERNS:
        matches = re.findall(pattern, text)
        placeholders.extend(matches)
    return list(set(placeholders))


def extract_placeholders_from_paragraph(paragraph) -> List[str]:
    full_text = ""
    for run in paragraph.runs:
        full_text += run.text
    return extract_placeholders_from_text(full_text)


def extract_placeholders_from_document(doc: Document) -> Dict[str, List[str]]:
    result = {
        "paragraphs": [],
        "tables": [],
        "headers": [],
        "footers": []
    }
    
    for paragraph in doc.paragraphs:
        placeholders = extract_placeholders_from_paragraph(paragraph)
        result["paragraphs"].extend(placeholders)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    placeholders = extract_placeholders_from_paragraph(paragraph)
                    result["tables"].extend(placeholders)
    
    for section in doc.sections:
        header = section.header
        for paragraph in header.paragraphs:
            placeholders = extract_placeholders_from_paragraph(paragraph)
            result["headers"].extend(placeholders)
        
        footer = section.footer
        for paragraph in footer.paragraphs:
            placeholders = extract_placeholders_from_paragraph(paragraph)
            result["footers"].extend(placeholders)
    
    for key in result:
        result[key] = list(set(result[key]))
    
    return result


def replace_in_runs(paragraph, replacements: Dict[str, str]):
    """Replace placeholders in paragraph runs while preserving formatting.
    
    Strategy:
    1. Try per-run replacement first (preserves all formatting)
    2. Fall back to cross-run replacement if placeholder spans multiple runs
    3. Final cleanup pass to remove any stray braces left behind
    """
    val_str = lambda v: str(v) if v is not None and v != "" else ""
    
    for run in paragraph.runs:
        for placeholder, value in replacements.items():
            patterns = [
                f"{{{{{placeholder}}}}}",
                f"{{{placeholder}}}",
                f"[[{placeholder}]]",
                f"[{placeholder}]",
                f"%{placeholder}%",
                f"<<{placeholder}>>",
                f"##{placeholder}##",
            ]
            for pattern in patterns:
                if pattern in run.text:
                    run.text = run.text.replace(pattern, val_str(value))
    
    full_text = "".join(run.text for run in paragraph.runs)
    
    replaced = False
    for placeholder, value in replacements.items():
        patterns = [
            f"{{{{{placeholder}}}}}",
            f"{{{placeholder}}}",
            f"[[{placeholder}]]",
            f"[{placeholder}]",
            f"%{placeholder}%",
            f"<<{placeholder}>>",
            f"##{placeholder}##",
        ]
        for pattern in patterns:
            if pattern in full_text:
                new_text = full_text.replace(pattern, val_str(value))
                if new_text != full_text:
                    full_text = new_text
                    replaced = True
    
    if replaced and paragraph.runs:
        paragraph.runs[0].text = full_text
        for i in range(1, len(paragraph.runs)):
            paragraph.runs[i].text = ""
    
    if replaced:
        for run in paragraph.runs:
            if run.text.strip() == '}' or run.text.strip() == '{':
                run.text = run.text.replace('}', '').replace('{', '')
                continue
            
            text = run.text
            has_valid_placeholder = bool(re.search(r'\{\{[^}]+\}\}|\{[^}]+\}', text))
            if not has_valid_placeholder:
                cleaned = re.sub(r'^\}+', '', text)
                cleaned = re.sub(r'\}+$', '', cleaned)
                cleaned = re.sub(r'^\{+', '', cleaned)
                cleaned = re.sub(r'\{+$', '', cleaned)
                if cleaned != text:
                    run.text = cleaned


def force_replace_remaining(paragraph, replacements: Dict[str, str]):
    """Force-replace any remaining placeholder patterns in paragraph text.
    This catches cases where run-level replacement missed due to complex run splitting."""
    full_text = "".join(run.text for run in paragraph.runs)
    original = full_text
    
    for placeholder, value in replacements.items():
        val = str(value) if value is not None and value != "" else ""
        patterns = [
            f"{{{{{placeholder}}}}}",
            f"{{{placeholder}}}",
            f"[[{placeholder}]]",
            f"[{placeholder}]",
            f"%{placeholder}%",
            f"<<{placeholder}>>",
            f"##{placeholder}##",
        ]
        for pattern in patterns:
            if pattern in full_text:
                full_text = full_text.replace(pattern, val)
    
    if full_text != original and paragraph.runs:
        paragraph.runs[0].text = full_text
        for i in range(1, len(paragraph.runs)):
            paragraph.runs[i].text = ""
        return True
    return False


def replace_placeholders_in_document(doc: Document, replacements: Dict[str, str]) -> int:
    count = 0
    
    all_paragraphs = []
    
    for paragraph in doc.paragraphs:
        all_paragraphs.append(paragraph)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    all_paragraphs.append(paragraph)
    
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            all_paragraphs.append(paragraph)
        for paragraph in section.footer.paragraphs:
            all_paragraphs.append(paragraph)
    
    for paragraph in all_paragraphs:
        original_text = "".join([run.text for run in paragraph.runs])
        replace_in_runs(paragraph, replacements)
        new_text = "".join([run.text for run in paragraph.runs])
        if original_text != new_text:
            count += 1
    
    forced = 0
    placeholder_pattern = re.compile(r'\{\{[^}]+\}\}|\{[^}]+\}|\[\[[^\]]+\]\]|\[[^\]]+\]|%[^%]+%|<<[^>]+>>|##[^#]+##')
    for paragraph in all_paragraphs:
        para_text = "".join(run.text for run in paragraph.runs)
        if placeholder_pattern.search(para_text):
            if force_replace_remaining(paragraph, replacements):
                forced += 1
    
    if forced > 0:
        print(f"[FORCE REPLACE] Fixed {forced} paragraphs with remaining placeholders")
    
    return count + forced


FALLBACK_FONT = "Roboto"
DEFAULT_TEXT_SIZE = Pt(10)
DEFAULT_TITLE_SIZE = Pt(12)


def apply_consistent_fonts(doc: Document):
    """Preserve original document fonts. Only apply fallback Google Font (Roboto) 
    when a run has no font set. Never override existing font names or sizes."""
    def ensure_font(run, is_title=False):
        if not run.font.name:
            run.font.name = FALLBACK_FONT
        if not run.font.size:
            run.font.size = DEFAULT_TITLE_SIZE if is_title else DEFAULT_TEXT_SIZE
    
    for paragraph in doc.paragraphs:
        is_title = paragraph.style and paragraph.style.name and ('Heading' in paragraph.style.name or 'Title' in paragraph.style.name)
        for run in paragraph.runs:
            ensure_font(run, is_title)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    is_title = paragraph.style and paragraph.style.name and ('Heading' in paragraph.style.name or 'Title' in paragraph.style.name)
                    for run in paragraph.runs:
                        ensure_font(run, is_title)
    
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            for run in paragraph.runs:
                ensure_font(run, is_title=False)
        for paragraph in section.footer.paragraphs:
            for run in paragraph.runs:
                ensure_font(run, is_title=False)


def load_template_metadata() -> Dict[str, Any]:
    metadata_file = os.path.join(SETTINGS_DIR, "templates_metadata.json")
    if os.path.exists(metadata_file):
        with open(metadata_file, "r") as f:
            return json.load(f)
    return {}


def save_template_metadata(metadata: Dict[str, Any]):
    metadata_file = os.path.join(SETTINGS_DIR, "templates_metadata.json")
    with open(metadata_file, "w") as f:
        json.dump(metadata, f, indent=2, default=str)


def get_current_date_str() -> str:
    """Get a random date within the last 14 days for 'current' dates."""
    days_ago = random.randint(0, 14)
    date = datetime.now() - timedelta(days=days_ago)
    return date.strftime("%Y-%m-%d")


def get_past_date_str() -> str:
    """Get a random date within the last 30 days for 'past' dates."""
    days_ago = random.randint(1, 30)
    date = datetime.now() - timedelta(days=days_ago)
    return date.strftime("%Y-%m-%d")


def get_future_date_str() -> str:
    """Get a random date within the next 30 days for 'future' dates."""
    days_ahead = random.randint(1, 30)
    date = datetime.now() + timedelta(days=days_ahead)
    return date.strftime("%Y-%m-%d")


def fetch_template_placeholders(template_name: str) -> List[Dict[str, Any]]:
    """Fetch placeholder mappings from Supabase document_template_fields table.
    
    CRITICAL: Query document_template_fields by template_file_name (exact match).
    
    Returns list of mappings with structure:
    - placeholder_name: str (the placeholder text)
    - source: 'database' or 'ai'
    - database_table: str (e.g., 'vessels', 'buyer_companies')
    - database_column: str (e.g., 'name', 'imo_number')
    """
    if not supabase:
        print("[MAPPINGS] Supabase not configured, returning empty mappings")
        return []
    
    try:
        file_name = template_name
        if not file_name.lower().endswith('.docx'):
            file_name = f"{file_name}.docx"
        
        print(f"[MAPPINGS] Querying document_template_fields by template_file_name: {file_name}")
        
        response = supabase.table("document_template_fields").select("*").eq("template_file_name", file_name).execute()
        
        if not response.data:
            template_base = template_name.replace('.docx', '').replace('.DOCX', '')
            print(f"[MAPPINGS] Exact match not found, trying with base name: {template_base}")
            response = supabase.table("document_template_fields").select("*").ilike("template_file_name", f"%{template_base}%").execute()
        
        if response.data:
            mappings = response.data
            db_count = len([m for m in mappings if m.get("source", "").lower() == "database"])
            ai_count = len([m for m in mappings if m.get("source", "").lower() == "ai"])
            print(f"[MAPPINGS] Found {len(mappings)} mappings: {db_count} database, {ai_count} ai")
            for m in mappings[:10]:
                print(f"[MAPPINGS]   - {m.get('placeholder_name')}: source={m.get('source')}, table={m.get('database_table')}, column={m.get('database_column')}")
            if len(mappings) > 10:
                print(f"[MAPPINGS]   ... and {len(mappings) - 10} more")
            return mappings
        else:
            print(f"[MAPPINGS] No mappings found in document_template_fields for: {file_name}")
            return []
            
    except Exception as e:
        print(f"[MAPPINGS] ERROR fetching document_template_fields: {e}")
        import traceback
        traceback.print_exc()
        return []


def fetch_value_from_database(table_name: str, field_name: str, entity_id: str = None) -> Optional[str]:
    """Fetch a specific field value from a database table."""
    if not supabase:
        return None
    
    try:
        if entity_id and is_valid_uuid(entity_id):
            response = supabase.table(table_name).select(field_name).eq("id", entity_id).limit(1).execute()
        else:
            response = supabase.table(table_name).select(field_name).limit(1).execute()
        
        if response.data and len(response.data) > 0:
            return response.data[0].get(field_name)
        return None
    except Exception as e:
        print(f"[DB FETCH] ERROR fetching {table_name}.{field_name}: {e}")
        return None


FIXED_PLACEHOLDER_VALUES = {
    "notary_name": "Valine Weisel",
    "witness_name_company": "Valine Weisel",
    "witness_name_authority": "Valine Weisel",
}

CURRENT_DATE_PLACEHOLDERS = [
    "effective_date", "issue_date", "certification_date", "notary_date",
    "approval_date", "document_date", "sgs_issue_date", "contract_date",
    "signing_date", "date", "commencement_date", "agreement_date",
    "execution_date", "issuance_date", "dated", "report_date",
    "inspection_date", "certificate_date", "authorization_date",
]

FUTURE_DATE_PLACEHOLDERS = [
    "amendment_issue_date", "future_issue_date", "revision_issue_date",
    "amendment_effective_date", "next_issuance_date", "compliance_date",
    "delivery_date", "eta", "expiry_date", "laycan_start", "laycan_end",
    "completion_date", "due_date", "renewal_date", "validity_date",
    "next_review_date", "deadline_date",
]

def _classify_date_placeholder(name: str) -> Optional[str]:
    """Return 'current' or 'future' if the placeholder looks like a date, else None.
    Checks FUTURE list first to avoid substring collisions (e.g. 'amendment_issue_date' contains 'issue_date')."""
    name_lower = name.lower().strip()
    for kw in FUTURE_DATE_PLACEHOLDERS:
        if kw in name_lower:
            return "future"
    for kw in CURRENT_DATE_PLACEHOLDERS:
        if kw in name_lower:
            return "current"
    if "date" in name_lower:
        return "current"
    return None


async def generate_ai_values(placeholders: List[str], context: Dict[str, str]) -> Dict[str, str]:
    if not openai_client or not placeholders:
        return {}
    
    today = datetime.now()
    today_str = today.strftime("%Y-%m-%d")
    
    past_days = random.randint(3, 21)
    current_date_obj = today - timedelta(days=past_days)
    current_date_formatted = current_date_obj.strftime("%B %d, %Y").replace(" 0", " ")
    
    future_days = random.randint(5, 30)
    future_date_obj = today + timedelta(days=future_days)
    future_date_formatted = future_date_obj.strftime("%B %d, %Y").replace(" 0", " ")
    
    pre_filled = {}
    remaining_placeholders = []
    
    for p in placeholders:
        p_lower = p.lower().strip()
        fixed_match = None
        for fixed_key, fixed_val in FIXED_PLACEHOLDER_VALUES.items():
            if fixed_key in p_lower or p_lower in fixed_key:
                fixed_match = fixed_val
                break
        
        if fixed_match:
            pre_filled[p] = fixed_match
            print(f"[FIXED VALUE] {p} = {fixed_match}")
            continue
        
        date_type = _classify_date_placeholder(p)
        if date_type == "current":
            pre_filled[p] = current_date_formatted
        elif date_type == "future":
            pre_filled[p] = future_date_formatted
        else:
            remaining_placeholders.append(p)
    
    if pre_filled:
        print(f"[AI DATES] Pre-filled {len(pre_filled)} placeholders: current={current_date_formatted}, future={future_date_formatted}")
    
    if not remaining_placeholders:
        return pre_filled
    
    context_str = "\n".join([f"- {k}: {v}" for k, v in context.items() if v])
    placeholders_str = "\n".join([f"- {p}" for p in remaining_placeholders])
    
    prompt = f"""Generate realistic values for these oil trading document placeholders:
{placeholders_str}

Context:
{context_str}

CRITICAL DATE RULES (MUST FOLLOW):
- Today's real date is: {today_str}
- If any placeholder looks like a date, use this EXACT value: "{current_date_formatted}"
- NEVER generate dates in YYYY-MM-DD format, always use human-readable like "{current_date_formatted}"
- NEVER use today's exact date

NUMERIC RESULT RULES:
- Placeholders starting with "Result" (like Result, Result1, Result2, Result3, etc.) are ALWAYS numeric values
- These represent oil trading calculation results (quantities, prices, totals, percentages)
- Generate realistic numeric values (integers or decimals) as strings, e.g. "1250000", "42.75", "0.85"
- Do NOT generate text descriptions for Result placeholders, only numbers

NAME GENERATION RULES (VERY IMPORTANT):
- NEVER use generic placeholder names like "John Smith", "Jane Doe", "John Doe", "James Smith", "Robert Johnson"
- Always generate realistic, UNIQUE, region-appropriate full names
- Use diverse names from different backgrounds (European, Middle Eastern, Asian, etc.)
- Names should sound like real professionals in oil trading/shipping/legal industries
- Examples of GOOD names: "Marcus Lindqvist", "Fatima Al-Rashidi", "Dimitri Volkov", "Helena Papadopoulos"
- Examples of BAD names: "John Smith", "Jane Doe", "Bob Jones", "Mary Johnson"

Return JSON only with placeholder names as keys and realistic values as strings.
Example: {{"contract_number": "OTC-2024-00123", "arbitration_clause": "ICC Rules of Arbitration"}}"""
    
    try:
        response = openai_client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "user", "content": prompt}],
            response_format={"type": "json_object"}
        )
        ai_result = json.loads(response.choices[0].message.content)
        for k, v in ai_result.items():
            if k not in pre_filled:
                pre_filled[k] = v
        return pre_filled
    except Exception as e:
        error_str = str(e).lower()
        if "quota" in error_str or "billing" in error_str or "insufficient" in error_str or "rate_limit" in error_str or "429" in error_str:
            print(f"[OPENAI QUOTA ERROR] {e}")
            raise HTTPException(
                status_code=402,
                detail={
                    "error": "openai_quota_exceeded",
                    "message": "OpenAI API quota exceeded or billing issue. Please recharge your OpenAI account immediately.",
                    "type": "billing_error"
                }
            )
        print(f"OpenAI API error: {e}")
        return pre_filled


@app.get("/")
async def root():
    return {
        "name": "PetroDealHub Document Processor API",
        "version": "1.0.0",
        "status": "running",
        "endpoints": {
            "health": "/health",
            "templates": "/templates",
            "upload_template": "/upload-template",
            "detect_placeholders": "/detect-placeholders/{template_name}",
            "process_document": "/process-document",
            "convert_to_pdf": "/convert-to-pdf",
            "database_tables": "/database-tables",
            "vessels": "/vessels",
            "plans": "/plans"
        }
    }


@app.get("/health")
async def health_check():
    return {"status": "healthy", "version": "1.0.0"}


@app.post("/upload-template")
async def upload_template(
    file: UploadFile = File(...),
    name: str = Form(...),
    description: Optional[str] = Form(None),
    font_family: Optional[str] = Form(None),
    font_size: Optional[int] = Form(None),
    plan_ids: Optional[str] = Form(None)
):
    if not file.filename.endswith(".docx"):
        raise HTTPException(status_code=400, detail="Only DOCX files are allowed")
    
    content = await file.read()
    template_id = str(uuid.uuid4())
    file_path = os.path.join(TEMPLATES_DIR, file.filename)
    
    async with aiofiles.open(file_path, "wb") as f:
        await f.write(content)
    
    doc = Document(BytesIO(content))
    placeholders_by_location = extract_placeholders_from_document(doc)
    all_placeholders = []
    for phs in placeholders_by_location.values():
        all_placeholders.extend(phs)
    all_placeholders = list(set(all_placeholders))
    
    plan_ids_list = json.loads(plan_ids) if plan_ids else ["basic", "professional", "enterprise"]
    
    metadata = load_template_metadata()
    metadata[template_id] = {
        "id": template_id,
        "name": name,
        "file_name": file.filename,
        "description": description,
        "placeholders": all_placeholders,
        "placeholder_count": len(all_placeholders),
        "font_family": font_family,
        "font_size": font_size,
        "file_size": len(content),
        "plan_ids": plan_ids_list,
        "is_active": True,
        "created_at": datetime.now().isoformat()
    }
    save_template_metadata(metadata)
    
    return {
        "success": True,
        "template_id": template_id,
        "file_name": file.filename,
        "placeholders": all_placeholders,
        "placeholder_count": len(all_placeholders),
        "file_size": len(content)
    }


@app.get("/templates")
async def list_templates():
    metadata = load_template_metadata()
    return {"templates": list(metadata.values())}


@app.get("/templates/{template_id}")
async def get_template(template_id: str):
    metadata = load_template_metadata()
    if template_id not in metadata:
        raise HTTPException(status_code=404, detail="Template not found")
    return metadata[template_id]


@app.delete("/templates/{template_name}")
async def delete_template(template_name: str):
    file_path = os.path.join(TEMPLATES_DIR, template_name)
    if os.path.exists(file_path):
        os.remove(file_path)
    
    metadata = load_template_metadata()
    for tid, tmpl in list(metadata.items()):
        if tmpl.get("file_name") == template_name:
            del metadata[tid]
            break
    save_template_metadata(metadata)
    
    return {"success": True}


@app.post("/templates/{template_id}/metadata")
async def update_template_metadata(template_id: str, body: TemplateMetadata):
    metadata = load_template_metadata()
    if template_id not in metadata:
        raise HTTPException(status_code=404, detail="Template not found")
    
    if body.display_name:
        metadata[template_id]["name"] = body.display_name
    if body.description is not None:
        metadata[template_id]["description"] = body.description
    if body.font_family:
        metadata[template_id]["font_family"] = body.font_family
    if body.font_size:
        metadata[template_id]["font_size"] = body.font_size
    if body.plan_ids:
        metadata[template_id]["plan_ids"] = body.plan_ids
    
    save_template_metadata(metadata)
    return {"success": True}


@app.get("/detect-placeholders/{template_name}")
async def detect_placeholders(template_name: str):
    file_path = os.path.join(TEMPLATES_DIR, template_name)
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="Template not found")
    
    doc = Document(file_path)
    placeholders_by_location = extract_placeholders_from_document(doc)
    all_placeholders = []
    for phs in placeholders_by_location.values():
        all_placeholders.extend(phs)
    all_placeholders = list(set(all_placeholders))
    
    return {
        "template_name": template_name,
        "placeholders": all_placeholders,
        "placeholder_count": len(all_placeholders),
        "by_location": placeholders_by_location
    }


@app.get("/database-tables")
async def get_database_tables():
    return {"tables": DATABASE_TABLES}


@app.get("/database-tables/{table_name}/columns")
async def get_table_columns(table_name: str):
    if not supabase:
        raise HTTPException(status_code=500, detail="Database not configured")
    
    if table_name not in DATABASE_TABLES:
        raise HTTPException(status_code=404, detail="Table not found")
    
    try:
        response = supabase.table(table_name).select("*").limit(1).execute()
        if response.data and len(response.data) > 0:
            sample = response.data[0]
            columns = []
            for key, value in sample.items():
                col_type = "text"
                if isinstance(value, int):
                    col_type = "integer"
                elif isinstance(value, float):
                    col_type = "numeric"
                elif isinstance(value, bool):
                    col_type = "boolean"
                columns.append({
                    "name": key,
                    "type": col_type,
                    "nullable": True
                })
            return {"table": table_name, "columns": columns}
        return {"table": table_name, "columns": []}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/placeholder-settings")
async def get_placeholder_settings(template_name: str = Query(...)):
    settings_file = os.path.join(SETTINGS_DIR, f"{template_name}.json")
    if os.path.exists(settings_file):
        with open(settings_file, "r") as f:
            data = json.load(f)
        return data
    
    metadata = load_template_metadata()
    template_id = None
    for tid, tmpl in metadata.items():
        if tmpl.get("file_name") == template_name:
            template_id = tid
            break
    
    return {
        "template_name": template_name,
        "template_id": template_id,
        "settings": {}
    }


@app.post("/placeholder-settings")
async def save_placeholder_settings(body: PlaceholderSettingsRequest):
    settings_file = os.path.join(SETTINGS_DIR, f"{body.template_name}.json")
    data = {
        "template_name": body.template_name,
        "template_id": body.template_id,
        "settings": body.settings
    }
    with open(settings_file, "w") as f:
        json.dump(data, f, indent=2)
    return {"success": True}


@app.get("/vessels")
async def list_vessels():
    if not supabase:
        raise HTTPException(status_code=500, detail="Database not configured")
    
    try:
        response = supabase.table("vessels").select("id, name, imo, vessel_type, flag").execute()
        return {"vessels": response.data}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/vessel/{identifier}")
async def get_vessel(identifier: str):
    if not supabase:
        raise HTTPException(status_code=500, detail="Database not configured")
    
    try:
        if identifier.isdigit():
            response = supabase.table("vessels").select("*").eq("id", int(identifier)).execute()
        else:
            response = supabase.table("vessels").select("*").eq("imo", identifier).execute()
        
        if response.data and len(response.data) > 0:
            return response.data[0]
        raise HTTPException(status_code=404, detail="Vessel not found")
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/process-document")
async def process_document(body: ProcessDocumentRequest):
    if not supabase:
        raise HTTPException(status_code=500, detail="Database not configured")
    
    template_name = body.template_name
    if not template_name.lower().endswith('.docx'):
        template_name = f"{template_name}.docx"
    
    file_path = os.path.join(TEMPLATES_DIR, template_name)
    if not os.path.exists(file_path):
        file_path_no_ext = os.path.join(TEMPLATES_DIR, body.template_name)
        if os.path.exists(file_path_no_ext):
            file_path = file_path_no_ext
            template_name = body.template_name
        else:
            raise HTTPException(status_code=404, detail=f"Template '{body.template_name}' not found")
    
    doc = Document(file_path)
    placeholders_by_location = extract_placeholders_from_document(doc)
    all_placeholders = set()
    for phs in placeholders_by_location.values():
        all_placeholders.update(phs)
    
    placeholder_data = {}
    from_database = 0
    from_ai = 0
    context = {}
    ai_only_placeholders = []
    
    print(f"\n[PROCESS] ========== TEMPLATE PLACEHOLDERS LOOKUP ==========")
    template_mappings = fetch_template_placeholders(template_name)
    
    if template_mappings:
        print(f"[PROCESS] Found {len(template_mappings)} mappings in document_template_fields table")
        
        print(f"[PROCESS] NOTE: database mappings will be resolved after vessel fetch (2nd pass)")
        
        for mapping in template_mappings:
            placeholder_name = mapping.get("placeholder_name", "")
            placeholder_name = re.sub(r'[\{\}\[\]<>%#]', '', placeholder_name).strip()
            source = mapping.get("source", "").lower()
            
            if source == "ai":
                ai_only_placeholders.append(placeholder_name)
                print(f"[PROCESS] MAPPING: {placeholder_name} <- AI (as configured)")
            elif source == "database":
                print(f"[PROCESS] MAPPING (queued): {placeholder_name} <- database (will resolve after entity fetch)")
    else:
        print("[PROCESS] No mappings found in document_template_fields, using fallback table fetches")
    
    print(f"[PROCESS] ============================================\n")
    
    vessel_buyer_company_uuid = None
    vessel_seller_company_uuid = None
    
    if body.vessel_id:
        try:
            response = supabase.table("vessels").select("*").eq("id", body.vessel_id).execute()
            if response.data:
                vessel = response.data[0]
                context["vessel_name"] = vessel.get("name", "")
                for key, value in vessel.items():
                    placeholder_data[f"vessel_{key}"] = value
                    placeholder_data[key] = value
                from_database += len(vessel)
                vessel_buyer_company_uuid = vessel.get("buyer_company_uuid")
                vessel_seller_company_uuid = vessel.get("seller_company_uuid")
                print(f"[PROCESS] Vessel FK: buyer_company_uuid={vessel_buyer_company_uuid}, seller_company_uuid={vessel_seller_company_uuid}")
        except Exception as e:
            print(f"Error fetching vessel: {e}")
    
    effective_buyer_id = body.buyer_id or vessel_buyer_company_uuid
    effective_seller_id = body.seller_id or vessel_seller_company_uuid
    
    if not effective_buyer_id:
        try:
            all_buyers = supabase.table("buyer_companies").select("id").execute()
            if all_buyers.data:
                random_buyer = random.choice(all_buyers.data)
                effective_buyer_id = random_buyer["id"]
                print(f"[PROCESS] Random buyer selected: {effective_buyer_id} (from {len(all_buyers.data)} available)")
        except Exception as e:
            print(f"[PROCESS] Error fetching random buyer: {e}")

    if not effective_seller_id:
        try:
            all_sellers = supabase.table("seller_companies").select("id").execute()
            if all_sellers.data:
                random_seller = random.choice(all_sellers.data)
                effective_seller_id = random_seller["id"]
                print(f"[PROCESS] Random seller selected: {effective_seller_id} (from {len(all_sellers.data)} available)")
        except Exception as e:
            print(f"[PROCESS] Error fetching random seller: {e}")

    if effective_buyer_id:
        buyer_id_str = str(effective_buyer_id)
        print(f"[PROCESS] Using buyer_id: {buyer_id_str} (from {'request' if body.buyer_id else 'vessel.buyer_company_uuid'})")
        try:
            response = supabase.table("buyer_companies").select("*").eq("id", buyer_id_str).execute()
            if response.data:
                buyer = response.data[0]
                context["buyer_name"] = buyer.get("name", "")
                for key, value in buyer.items():
                    placeholder_data[f"buyer_{key}"] = value
                from_database += len(buyer)
        except Exception as e:
            print(f"Error fetching buyer: {e}")
    
    if effective_seller_id:
        seller_id_str = str(effective_seller_id)
        print(f"[PROCESS] Using seller_id: {seller_id_str} (from {'request' if body.seller_id else 'vessel.seller_company_uuid'})")
        try:
            response = supabase.table("seller_companies").select("*").eq("id", seller_id_str).execute()
            if response.data:
                seller = response.data[0]
                context["seller_name"] = seller.get("name", "")
                for key, value in seller.items():
                    placeholder_data[f"seller_{key}"] = value
                from_database += len(seller)
        except Exception as e:
            print(f"Error fetching seller: {e}")
    
    if body.product_id:
        try:
            response = supabase.table("oil_products").select("*").eq("id", body.product_id).execute()
            if response.data:
                product = response.data[0]
                context["cargo_type"] = product.get("commodity_name", "")
                for key, value in product.items():
                    placeholder_data[f"product_{key}"] = value
                from_database += len(product)
        except Exception as e:
            print(f"Error fetching product: {e}")
    
    if body.refinery_id:
        try:
            response = supabase.table("refineries").select("*").eq("id", body.refinery_id).execute()
            if response.data:
                refinery = response.data[0]
                for key, value in refinery.items():
                    placeholder_data[f"refinery_{key}"] = value
                from_database += len(refinery)
        except Exception as e:
            print(f"Error fetching refinery: {e}")
    
    if body.departure_port_id:
        try:
            response = supabase.table("ports").select("*").eq("id", body.departure_port_id).execute()
            if response.data:
                port = response.data[0]
                context["departure_port"] = port.get("name", "")
                for key, value in port.items():
                    placeholder_data[f"departure_port_{key}"] = value
                from_database += len(port)
        except Exception as e:
            print(f"Error fetching departure port: {e}")
    
    if body.destination_port_id:
        try:
            response = supabase.table("ports").select("*").eq("id", body.destination_port_id).execute()
            if response.data:
                port = response.data[0]
                context["destination_port"] = port.get("name", "")
                for key, value in port.items():
                    placeholder_data[f"destination_port_{key}"] = value
                from_database += len(port)
        except Exception as e:
            print(f"Error fetching destination port: {e}")
    
    if body.buyer_bank_id:
        try:
            response = supabase.table("buyer_company_bank_accounts").select("*").eq("id", body.buyer_bank_id).execute()
            if response.data:
                bank = response.data[0]
                for key, value in bank.items():
                    placeholder_data[f"buyer_bank_{key}"] = value
                from_database += len(bank)
        except Exception as e:
            print(f"Error fetching buyer bank: {e}")
    
    if body.seller_bank_id:
        try:
            response = supabase.table("seller_company_bank_accounts").select("*").eq("id", body.seller_bank_id).execute()
            if response.data:
                bank = response.data[0]
                for key, value in bank.items():
                    placeholder_data[f"seller_bank_{key}"] = value
                from_database += len(bank)
        except Exception as e:
            print(f"Error fetching seller bank: {e}")
    
    if template_mappings:
        print(f"\n[PROCESS] ========== RESOLVING TEMPLATE MAPPINGS (2nd pass) ==========")
        entity_ids = {
            "vessels": body.vessel_id,
            "buyer_companies": effective_buyer_id,
            "seller_companies": effective_seller_id,
            "oil_products": body.product_id,
            "refineries": body.refinery_id,
            "ports": body.departure_port_id or body.destination_port_id,
            "buyer_company_bank_accounts": body.buyer_bank_id,
            "seller_company_bank_accounts": body.seller_bank_id
        }
        print(f"[PROCESS] Entity IDs: buyer={effective_buyer_id} (from {'request' if body.buyer_id else 'vessel'}), seller={effective_seller_id} (from {'request' if body.seller_id else 'vessel'})")
        
        for mapping in template_mappings:
            placeholder_name = mapping.get("placeholder_name", "")
            placeholder_name = re.sub(r'[\{\}\[\]<>%#]', '', placeholder_name).strip()
            source = mapping.get("source", "").lower()
            db_table = mapping.get("database_table", "")
            db_column = mapping.get("database_column", "")
            
            if source == "database" and db_table and db_column:
                entity_id = entity_ids.get(db_table)
                value = fetch_value_from_database(db_table, db_column, entity_id)
                if value is not None:
                    was_set = placeholder_name in placeholder_data
                    placeholder_data[placeholder_name] = value
                    from_database += 1
                    if was_set:
                        print(f"[PROCESS] OVERWRITE: {placeholder_name} <- {db_table}.{db_column} = '{str(value)[:50]}...' (mapping overrides prefix)")
                    else:
                        print(f"[PROCESS] RESOLVED: {placeholder_name} <- {db_table}.{db_column} = '{str(value)[:50]}...'")
                else:
                    print(f"[PROCESS] NOT FOUND: {placeholder_name} <- {db_table}.{db_column} (entity_id={entity_id}), will use AI")
                    ai_only_placeholders.append(placeholder_name)
        print(f"[PROCESS] ============================================================\n")
    
    missing_placeholders = []
    for ph in all_placeholders:
        found = False
        for key in placeholder_data:
            if ph.lower() == key.lower() or ph.lower().replace("_", "") == key.lower().replace("_", ""):
                found = True
                break
        if not found:
            missing_placeholders.append(ph)
    
    placeholders_for_ai = list(set(missing_placeholders + ai_only_placeholders))
    print(f"[PROCESS] Placeholders for AI generation: {len(placeholders_for_ai)} (missing: {len(missing_placeholders)}, mapped to AI: {len(ai_only_placeholders)})")
    
    ai_generated_placeholders = []
    if placeholders_for_ai and openai_client:
        ai_values = await generate_ai_values(placeholders_for_ai, context)
        for ph, value in ai_values.items():
            placeholder_data[ph] = value
            ai_generated_placeholders.append(ph)
            from_ai += 1
        missing_placeholders = [p for p in missing_placeholders if p not in ai_values]
        print(f"[PROCESS] AI-generated placeholders: {ai_generated_placeholders[:20]}...")
    
    print(f"\n[PROCESS] ========== FINAL SUMMARY ==========")
    print(f"[PROCESS] From Database: {from_database} values")
    print(f"[PROCESS] From AI: {from_ai} values")
    print(f"[PROCESS] Still missing: {len(missing_placeholders)}")
    print(f"[PROCESS] =====================================\n")
    
    doc = Document(file_path)
    replacements_made = replace_placeholders_in_document(doc, placeholder_data)
    apply_consistent_fonts(doc)
    
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    docx_bytes = output.read()
    
    pdf_bytes, conversion_method = convert_docx_to_image_pdf(docx_bytes, dpi=200)
    pdf_base64 = base64.b64encode(pdf_bytes).decode("utf-8")
    
    base_name = template_name.replace(".docx", "").replace(".DOCX", "")
    
    return {
        "success": True,
        "pdf_base64": pdf_base64,
        "pdf_file_name": f"{base_name}_filled.pdf",
        "replacements_made": replacements_made,
        "from_database": from_database,
        "from_ai": from_ai,
        "ai_generated_placeholders": ai_generated_placeholders,
        "conversion_method": conversion_method,
        "missing_placeholders": missing_placeholders
    }


@app.post("/generate-document")
async def generate_document(body: GenerateDocumentRequest):
    """Generate a document with vessel_imo support - alias for process-document"""
    if not supabase:
        raise HTTPException(status_code=500, detail="Database not configured")
    
    template_name = body.template_name
    if not template_name.lower().endswith('.docx'):
        template_name = f"{template_name}.docx"
    
    file_path = os.path.join(TEMPLATES_DIR, template_name)
    if not os.path.exists(file_path):
        file_path_no_ext = os.path.join(TEMPLATES_DIR, body.template_name)
        if os.path.exists(file_path_no_ext):
            file_path = file_path_no_ext
            template_name = body.template_name
        else:
            raise HTTPException(status_code=404, detail=f"Template '{body.template_name}' not found")
    
    print(f"\n[GENERATE] ========== REQUEST PARAMETERS ==========")
    print(f"[GENERATE] template_name: {body.template_name}")
    print(f"[GENERATE] vessel_imo: {body.vessel_imo}")
    print(f"[GENERATE] vessel_id: {body.vessel_id}")
    print(f"[GENERATE] buyer_id: {body.buyer_id} (type: {type(body.buyer_id).__name__})")
    print(f"[GENERATE] seller_id: {body.seller_id} (type: {type(body.seller_id).__name__})")
    print(f"[GENERATE] product_id: {body.product_id}")
    print(f"[GENERATE] refinery_id: {body.refinery_id}")
    print(f"[GENERATE] departure_port_id: {body.departure_port_id}")
    print(f"[GENERATE] destination_port_id: {body.destination_port_id}")
    print(f"[GENERATE] buyer_bank_id: {body.buyer_bank_id}")
    print(f"[GENERATE] seller_bank_id: {body.seller_bank_id}")
    print(f"[GENERATE] ================================================\n")
    
    vessel_id = body.vessel_id
    if body.vessel_imo and not vessel_id:
        try:
            response = supabase.table("vessels").select("id").eq("imo", body.vessel_imo).execute()
            if response.data and len(response.data) > 0:
                vessel_id = response.data[0]["id"]
                print(f"[GENERATE] Resolved vessel_imo {body.vessel_imo} -> vessel_id {vessel_id}")
            else:
                raise HTTPException(status_code=404, detail=f"Vessel with IMO '{body.vessel_imo}' not found")
        except HTTPException:
            raise
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Error looking up vessel: {str(e)}")
    
    doc = Document(file_path)
    placeholders_by_location = extract_placeholders_from_document(doc)
    all_placeholders = set()
    for phs in placeholders_by_location.values():
        all_placeholders.update(phs)
    
    placeholder_data = {}
    database_sources = {}
    from_database = 0
    from_ai = 0
    context = {}
    ai_only_placeholders = []
    
    print(f"\n[GENERATE] ========== TEMPLATE PLACEHOLDERS LOOKUP ==========")
    template_mappings = fetch_template_placeholders(template_name)
    
    if template_mappings:
        print(f"[GENERATE] Found {len(template_mappings)} mappings in document_template_fields table")
        
        print(f"[GENERATE] NOTE: entity_ids will be finalized after vessel lookup (buyer/seller may come from vessel)")
        
        for mapping in template_mappings:
            placeholder_name = mapping.get("placeholder_name", "")
            placeholder_name = re.sub(r'[\{\}\[\]<>%#]', '', placeholder_name).strip()
            source = mapping.get("source", "").lower()
            db_table = mapping.get("database_table", "")
            db_column = mapping.get("database_column", "")
            
            if source == "database" and db_table and db_column:
                ai_only_placeholders_temp = []
                print(f"[GENERATE] MAPPING (queued): {placeholder_name} <- {db_table}.{db_column} (will resolve after entity fetch)")
            elif source == "ai":
                ai_only_placeholders.append(placeholder_name)
                print(f"[GENERATE] MAPPING: {placeholder_name} <- AI (as configured)")
            else:
                print(f"[GENERATE] MAPPING: {placeholder_name} <- Unknown source: {source}")
    else:
        print("[GENERATE] No mappings found in document_template_fields, using fallback table fetches")
    
    print(f"[GENERATE] ============================================\n")
    
    vessel_buyer_company_uuid = None
    vessel_seller_company_uuid = None
    
    if vessel_id:
        try:
            response = supabase.table("vessels").select("*").eq("id", vessel_id).execute()
            if response.data:
                vessel = response.data[0]
                context["vessel_name"] = vessel.get("name", "")
                print(f"[GENERATE] VESSELS table: fetched {len(vessel)} fields for id={vessel_id}")
                for key, value in vessel.items():
                    placeholder_data[f"vessel_{key}"] = value
                    placeholder_data[key] = value
                    database_sources[f"vessel_{key}"] = "vessels"
                    database_sources[key] = "vessels"
                from_database += len(vessel)
                
                vessel_buyer_company_uuid = vessel.get("buyer_company_uuid")
                vessel_seller_company_uuid = vessel.get("seller_company_uuid")
                print(f"[GENERATE] Vessel FK: buyer_company_uuid={vessel_buyer_company_uuid}, seller_company_uuid={vessel_seller_company_uuid}")
            else:
                print(f"[GENERATE] WARNING: No data found in vessels for id={vessel_id}")
        except Exception as e:
            print(f"[GENERATE] ERROR fetching vessel: {e}")
    
    effective_buyer_id = body.buyer_id or vessel_buyer_company_uuid
    effective_seller_id = body.seller_id or vessel_seller_company_uuid
    
    if not effective_buyer_id:
        try:
            all_buyers = supabase.table("buyer_companies").select("id").execute()
            if all_buyers.data:
                random_buyer = random.choice(all_buyers.data)
                effective_buyer_id = random_buyer["id"]
                print(f"[GENERATE] Random buyer selected: {effective_buyer_id} (from {len(all_buyers.data)} available)")
        except Exception as e:
            print(f"[GENERATE] Error fetching random buyer: {e}")

    if not effective_seller_id:
        try:
            all_sellers = supabase.table("seller_companies").select("id").execute()
            if all_sellers.data:
                random_seller = random.choice(all_sellers.data)
                effective_seller_id = random_seller["id"]
                print(f"[GENERATE] Random seller selected: {effective_seller_id} (from {len(all_sellers.data)} available)")
        except Exception as e:
            print(f"[GENERATE] Error fetching random seller: {e}")

    if effective_buyer_id:
        buyer_id_str = str(effective_buyer_id)
        print(f"[GENERATE] Using buyer_id: {buyer_id_str} (from {'request' if body.buyer_id else 'vessel.buyer_company_uuid'})")
        try:
            response = supabase.table("buyer_companies").select("*").eq("id", buyer_id_str).execute()
            if response.data:
                buyer = response.data[0]
                context["buyer_name"] = buyer.get("name", "")
                print(f"[GENERATE] BUYER_COMPANIES table: fetched {len(buyer)} fields for id={buyer_id_str}")
                for key, value in buyer.items():
                    placeholder_data[f"buyer_{key}"] = value
                    database_sources[f"buyer_{key}"] = "buyer_companies"
                from_database += len(buyer)
            else:
                print(f"[GENERATE] WARNING: No data found in buyer_companies for id={buyer_id_str}")
        except Exception as e:
            print(f"[GENERATE] ERROR fetching buyer: {e}")
    
    if effective_seller_id:
        seller_id_str = str(effective_seller_id)
        print(f"[GENERATE] Using seller_id: {seller_id_str} (from {'request' if body.seller_id else 'vessel.seller_company_uuid'})")
        try:
            response = supabase.table("seller_companies").select("*").eq("id", seller_id_str).execute()
            if response.data:
                seller = response.data[0]
                context["seller_name"] = seller.get("name", "")
                print(f"[GENERATE] SELLER_COMPANIES table: fetched {len(seller)} fields for id={seller_id_str}")
                for key, value in seller.items():
                    placeholder_data[f"seller_{key}"] = value
                    database_sources[f"seller_{key}"] = "seller_companies"
                from_database += len(seller)
            else:
                print(f"[GENERATE] WARNING: No data found in seller_companies for id={seller_id_str}")
        except Exception as e:
            print(f"[GENERATE] ERROR fetching seller: {e}")
    
    if body.product_id:
        if not is_valid_uuid(body.product_id):
            print(f"[GENERATE] ERROR: product_id '{body.product_id}' is not a valid UUID format")
            raise HTTPException(status_code=400, detail=f"product_id '{body.product_id}' is not a valid UUID. Please provide a valid UUID from oil_products table.")
        try:
            print(f"[GENERATE] Querying oil_products with id='{body.product_id}'")
            response = supabase.table("oil_products").select("*").eq("id", body.product_id).execute()
            if response.data:
                product = response.data[0]
                context["cargo_type"] = product.get("commodity_name", "")
                print(f"[GENERATE] OIL_PRODUCTS table: fetched {len(product)} fields for id={body.product_id}")
                for key, value in product.items():
                    placeholder_data[f"product_{key}"] = value
                    database_sources[f"product_{key}"] = "oil_products"
                from_database += len(product)
            else:
                print(f"[GENERATE] WARNING: No data found in oil_products for id={body.product_id}")
        except Exception as e:
            print(f"[GENERATE] ERROR fetching product: {e}")
    
    if body.refinery_id:
        if not is_valid_uuid(body.refinery_id):
            print(f"[GENERATE] ERROR: refinery_id '{body.refinery_id}' is not a valid UUID format")
            raise HTTPException(status_code=400, detail=f"refinery_id '{body.refinery_id}' is not a valid UUID. Please provide a valid UUID from refineries table.")
        try:
            print(f"[GENERATE] Querying refineries with id='{body.refinery_id}'")
            response = supabase.table("refineries").select("*").eq("id", body.refinery_id).execute()
            if response.data:
                refinery = response.data[0]
                print(f"[GENERATE] REFINERIES table: fetched {len(refinery)} fields for id={body.refinery_id}")
                for key, value in refinery.items():
                    placeholder_data[f"refinery_{key}"] = value
                    database_sources[f"refinery_{key}"] = "refineries"
                from_database += len(refinery)
            else:
                print(f"[GENERATE] WARNING: No data found in refineries for id={body.refinery_id}")
        except Exception as e:
            print(f"[GENERATE] ERROR fetching refinery: {e}")
    
    if body.departure_port_id:
        try:
            response = supabase.table("ports").select("*").eq("id", body.departure_port_id).execute()
            if response.data:
                port = response.data[0]
                context["departure_port"] = port.get("name", "")
                print(f"[GENERATE] PORTS table (departure): fetched {len(port)} fields for id={body.departure_port_id}")
                for key, value in port.items():
                    placeholder_data[f"departure_port_{key}"] = value
                    database_sources[f"departure_port_{key}"] = "ports"
                from_database += len(port)
            else:
                print(f"[GENERATE] WARNING: No data found in ports for departure id={body.departure_port_id}")
        except Exception as e:
            print(f"[GENERATE] ERROR fetching departure port: {e}")
    
    if body.destination_port_id:
        try:
            response = supabase.table("ports").select("*").eq("id", body.destination_port_id).execute()
            if response.data:
                port = response.data[0]
                context["destination_port"] = port.get("name", "")
                print(f"[GENERATE] PORTS table (destination): fetched {len(port)} fields for id={body.destination_port_id}")
                for key, value in port.items():
                    placeholder_data[f"destination_port_{key}"] = value
                    database_sources[f"destination_port_{key}"] = "ports"
                from_database += len(port)
            else:
                print(f"[GENERATE] WARNING: No data found in ports for destination id={body.destination_port_id}")
        except Exception as e:
            print(f"[GENERATE] ERROR fetching destination port: {e}")
    
    if body.buyer_bank_id:
        try:
            print(f"[GENERATE] Querying buyer_company_bank_accounts with id='{body.buyer_bank_id}'")
            response = supabase.table("buyer_company_bank_accounts").select("*").eq("id", body.buyer_bank_id).execute()
            if response.data:
                bank = response.data[0]
                print(f"[GENERATE] BUYER_BANK table: fetched {len(bank)} fields for id={body.buyer_bank_id}")
                for key, value in bank.items():
                    placeholder_data[f"buyer_bank_{key}"] = value
                    database_sources[f"buyer_bank_{key}"] = "buyer_company_bank_accounts"
                from_database += len(bank)
            else:
                print(f"[GENERATE] WARNING: No data found in buyer_company_bank_accounts for id={body.buyer_bank_id}")
        except Exception as e:
            print(f"[GENERATE] ERROR fetching buyer bank: {e}")
    elif effective_buyer_id:
        try:
            print(f"[GENERATE] Auto-fetching buyer bank by company_id='{effective_buyer_id}'")
            response = supabase.table("buyer_company_bank_accounts").select("*").eq("company_id", str(effective_buyer_id)).execute()
            if response.data:
                bank = response.data[0]
                print(f"[GENERATE] BUYER_BANK table (auto): fetched {len(bank)} fields for company_id={effective_buyer_id}")
                for key, value in bank.items():
                    placeholder_data[f"buyer_bank_{key}"] = value
                    database_sources[f"buyer_bank_{key}"] = "buyer_company_bank_accounts"
                from_database += len(bank)
            else:
                print(f"[GENERATE] No buyer bank found for company_id={effective_buyer_id}")
        except Exception as e:
            print(f"[GENERATE] ERROR auto-fetching buyer bank: {e}")
    
    if body.seller_bank_id:
        try:
            print(f"[GENERATE] Querying seller_company_bank_accounts with id='{body.seller_bank_id}'")
            response = supabase.table("seller_company_bank_accounts").select("*").eq("id", body.seller_bank_id).execute()
            if response.data:
                bank = response.data[0]
                print(f"[GENERATE] SELLER_BANK table: fetched {len(bank)} fields for id={body.seller_bank_id}")
                for key, value in bank.items():
                    placeholder_data[f"seller_bank_{key}"] = value
                    database_sources[f"seller_bank_{key}"] = "seller_company_bank_accounts"
                from_database += len(bank)
            else:
                print(f"[GENERATE] WARNING: No data found in seller_company_bank_accounts for id={body.seller_bank_id}")
        except Exception as e:
            print(f"[GENERATE] ERROR fetching seller bank: {e}")
    elif effective_seller_id:
        try:
            print(f"[GENERATE] Auto-fetching seller bank by company_id='{effective_seller_id}'")
            response = supabase.table("seller_company_bank_accounts").select("*").eq("company_id", str(effective_seller_id)).execute()
            if response.data:
                bank = response.data[0]
                print(f"[GENERATE] SELLER_BANK table (auto): fetched {len(bank)} fields for company_id={effective_seller_id}")
                for key, value in bank.items():
                    placeholder_data[f"seller_bank_{key}"] = value
                    database_sources[f"seller_bank_{key}"] = "seller_company_bank_accounts"
                from_database += len(bank)
            else:
                print(f"[GENERATE] No seller bank found for company_id={effective_seller_id}")
        except Exception as e:
            print(f"[GENERATE] ERROR auto-fetching seller bank: {e}")
    
    if template_mappings:
        print(f"\n[GENERATE] ========== RESOLVING TEMPLATE MAPPINGS (2nd pass) ==========")
        entity_ids = {
            "vessels": vessel_id,
            "buyer_companies": effective_buyer_id,
            "seller_companies": effective_seller_id,
            "oil_products": body.product_id,
            "refineries": body.refinery_id,
            "ports": body.departure_port_id or body.destination_port_id,
            "buyer_company_bank_accounts": body.buyer_bank_id,
            "seller_company_bank_accounts": body.seller_bank_id
        }
        print(f"[GENERATE] Entity IDs: buyer={effective_buyer_id} (from {'request' if body.buyer_id else 'vessel'}), seller={effective_seller_id} (from {'request' if body.seller_id else 'vessel'})")
        
        for mapping in template_mappings:
            placeholder_name = mapping.get("placeholder_name", "")
            placeholder_name = re.sub(r'[\{\}\[\]<>%#]', '', placeholder_name).strip()
            source = mapping.get("source", "").lower()
            db_table = mapping.get("database_table", "")
            db_column = mapping.get("database_column", "")
            
            if source == "database" and db_table and db_column:
                entity_id = entity_ids.get(db_table)
                value = fetch_value_from_database(db_table, db_column, entity_id)
                if value is not None:
                    was_set = placeholder_name in placeholder_data
                    placeholder_data[placeholder_name] = value
                    database_sources[placeholder_name] = f"{db_table}.{db_column}"
                    from_database += 1
                    if was_set:
                        print(f"[GENERATE] OVERWRITE: {placeholder_name} <- {db_table}.{db_column} = '{str(value)[:50]}...' (mapping overrides prefix)")
                    else:
                        print(f"[GENERATE] RESOLVED: {placeholder_name} <- {db_table}.{db_column} = '{str(value)[:50]}...'")
                else:
                    print(f"[GENERATE] NOT FOUND: {placeholder_name} <- {db_table}.{db_column} (entity_id={entity_id}), will use AI")
                    ai_only_placeholders.append(placeholder_name)
        print(f"[GENERATE] ============================================================\n")
    
    print(f"\n[GENERATE] ========== DATABASE SUMMARY ==========")
    print(f"[GENERATE] Total placeholders from database: {from_database}")
    print(f"[GENERATE] Placeholder data keys: {list(placeholder_data.keys())[:20]}...")
    print(f"[GENERATE] ==========================================\n")
    
    missing_placeholders = []
    for ph in all_placeholders:
        found = False
        for key in placeholder_data:
            if ph.lower() == key.lower() or ph.lower().replace("_", "") == key.lower().replace("_", ""):
                found = True
                break
        if not found:
            missing_placeholders.append(ph)
    
    placeholders_for_ai = list(set(missing_placeholders + ai_only_placeholders))
    print(f"[GENERATE] Placeholders for AI generation: {len(placeholders_for_ai)} (missing: {len(missing_placeholders)}, mapped to AI: {len(ai_only_placeholders)})")
    
    ai_generated_placeholders = []
    if placeholders_for_ai and openai_client:
        ai_values = await generate_ai_values(placeholders_for_ai, context)
        for ph, value in ai_values.items():
            placeholder_data[ph] = value
            ai_generated_placeholders.append(ph)
            from_ai += 1
        missing_placeholders = [p for p in missing_placeholders if p not in ai_values]
        print(f"[GENERATE] AI-generated placeholders: {ai_generated_placeholders[:20]}...")
    
    print(f"\n[GENERATE] ========== FINAL SUMMARY ==========")
    print(f"[GENERATE] From Database: {from_database} values")
    print(f"[GENERATE] From AI: {from_ai} values")
    print(f"[GENERATE] Still missing: {len(missing_placeholders)}")
    print(f"[GENERATE] =====================================\n")
    
    doc = Document(file_path)
    replacements_made = replace_placeholders_in_document(doc, placeholder_data)
    apply_consistent_fonts(doc)
    
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    docx_bytes = output.read()
    
    pdf_bytes, conversion_method = convert_docx_to_image_pdf(docx_bytes, dpi=200)
    pdf_base64 = base64.b64encode(pdf_bytes).decode("utf-8")
    
    base_name = template_name.replace(".docx", "").replace(".DOCX", "")
    
    return {
        "success": True,
        "pdf_base64": pdf_base64,
        "pdf_file_name": f"{base_name}_filled.pdf",
        "replacements_made": replacements_made,
        "from_database": from_database,
        "from_ai": from_ai,
        "ai_generated_placeholders": ai_generated_placeholders,
        "conversion_method": conversion_method,
        "missing_placeholders": missing_placeholders
    }


@app.post("/upload-csv")
async def upload_csv(
    file: UploadFile = File(...),
    data_type: str = Form(...)
):
    if not file.filename.endswith(".csv"):
        raise HTTPException(status_code=400, detail="Only CSV files are allowed")
    
    content = await file.read()
    csv_id = str(uuid.uuid4())
    file_path = os.path.join(CSV_DIR, file.filename)
    
    async with aiofiles.open(file_path, "wb") as f:
        await f.write(content)
    
    df = pd.read_csv(BytesIO(content))
    headers = list(df.columns)
    row_count = len(df)
    
    csv_metadata_file = os.path.join(SETTINGS_DIR, "csv_metadata.json")
    csv_metadata = {}
    if os.path.exists(csv_metadata_file):
        with open(csv_metadata_file, "r") as f:
            csv_metadata = json.load(f)
    
    csv_metadata[csv_id] = {
        "id": csv_id,
        "name": file.filename,
        "data_type": data_type,
        "headers": headers,
        "row_count": row_count,
        "created_at": datetime.now().isoformat()
    }
    
    with open(csv_metadata_file, "w") as f:
        json.dump(csv_metadata, f, indent=2)
    
    return {
        "success": True,
        "csv_id": csv_id,
        "file_name": file.filename,
        "headers": headers,
        "row_count": row_count
    }


@app.get("/csv-files")
async def list_csv_files():
    csv_metadata_file = os.path.join(SETTINGS_DIR, "csv_metadata.json")
    if os.path.exists(csv_metadata_file):
        with open(csv_metadata_file, "r") as f:
            csv_metadata = json.load(f)
        return {"csv_files": list(csv_metadata.values())}
    return {"csv_files": []}


@app.get("/data/all")
async def get_all_data_sources():
    data_sources = {}
    
    for table in DATABASE_TABLES:
        data_sources[table] = {
            "type": "database",
            "name": table.replace("_", " ").title(),
            "row_count": 0
        }
        if supabase:
            try:
                response = supabase.table(table).select("id", count="exact").execute()
                data_sources[table]["row_count"] = response.count or 0
            except:
                pass
    
    csv_metadata_file = os.path.join(SETTINGS_DIR, "csv_metadata.json")
    if os.path.exists(csv_metadata_file):
        with open(csv_metadata_file, "r") as f:
            csv_metadata = json.load(f)
        for csv_id, csv_data in csv_metadata.items():
            data_sources[f"csv_{csv_data['data_type']}"] = {
                "type": "csv",
                "name": csv_data["name"],
                "row_count": csv_data["row_count"]
            }
    
    return {"data_sources": data_sources}


DEFAULT_PLANS = {
    "basic": {
        "id": "basic",
        "plan_name": "Basic",
        "plan_tier": "basic",
        "can_download": True,
        "max_downloads_per_month": 10,
        "allowed_templates": ["vessel_certificate"]
    },
    "professional": {
        "id": "professional",
        "plan_name": "Professional",
        "plan_tier": "professional",
        "can_download": True,
        "max_downloads_per_month": 50,
        "allowed_templates": ["vessel_certificate", "sales_purchase_agreement", "bill_of_lading"]
    },
    "enterprise": {
        "id": "enterprise",
        "plan_name": "Enterprise",
        "plan_tier": "enterprise",
        "can_download": True,
        "max_downloads_per_month": -1,
        "allowed_templates": ["*"]
    }
}


@app.get("/plans")
async def get_plans():
    return {"plans": DEFAULT_PLANS}


@app.get("/plans-db")
async def get_plans_from_db():
    if not supabase:
        return {"plans": DEFAULT_PLANS}
    
    try:
        response = supabase.table("subscription_plans").select("*").execute()
        if response.data:
            plans = {}
            for plan in response.data:
                plan_id = plan.get("id", plan.get("plan_id", str(uuid.uuid4())))
                plans[plan_id] = plan
            return {"plans": plans}
        return {"plans": DEFAULT_PLANS}
    except Exception as e:
        return {"plans": DEFAULT_PLANS}


@app.post("/update-plan")
async def update_plan(body: PlanUpdateRequest):
    if not supabase:
        raise HTTPException(status_code=500, detail="Database not configured")
    
    try:
        response = supabase.table("subscription_plans").update(body.plan_data).eq("id", body.plan_id).execute()
        return {"success": True}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


def convert_docx_to_pdf_direct(docx_bytes: bytes) -> bytes:
    """Convert DOCX to PDF directly using LibreOffice (faster, text selectable)"""
    with tempfile.TemporaryDirectory() as temp_dir:
        docx_path = os.path.join(temp_dir, "document.docx")
        with open(docx_path, "wb") as f:
            f.write(docx_bytes)
        
        result = subprocess.run([
            "libreoffice", "--headless", "--convert-to", "pdf",
            "--outdir", temp_dir, docx_path
        ], capture_output=True, timeout=120)
        
        if result.returncode != 0:
            raise Exception(f"LibreOffice conversion failed: {result.stderr.decode()}")
        
        pdf_path = os.path.join(temp_dir, "document.pdf")
        if not os.path.exists(pdf_path):
            raise Exception("PDF file was not created by LibreOffice")
        
        with open(pdf_path, "rb") as f:
            return f.read()


def convert_docx_to_image_pdf(docx_bytes: bytes, dpi: int = 200) -> tuple:
    """Convert DOCX to image-based PDF (non-editable) using hybrid method:
    
    HYBRID PIPELINE:
    1. Try docx2pdf first (if available)
    2. If docx2pdf fails, fallback to LibreOffice
    3. PDF -> Images (pdf2image/poppler) at specified DPI
    4. Images -> PDF (img2pdf)
    
    Args:
        docx_bytes: The DOCX file as bytes
        dpi: Resolution for image conversion (default 200 for better quality)
    
    Returns:
        tuple: (PDF bytes, conversion_method used)
    """
    import logging
    logger = logging.getLogger(__name__)
    
    conversion_method = "unknown"
    
    with tempfile.TemporaryDirectory() as temp_dir:
        docx_path = os.path.join(temp_dir, "document.docx")
        pdf_path = os.path.join(temp_dir, "document.pdf")
        
        with open(docx_path, "wb") as f:
            f.write(docx_bytes)
        logger.info(f"Step 1: Saved DOCX to {docx_path}, size: {len(docx_bytes)} bytes")
        print(f"[PDF] Step 1: Saved DOCX, size: {len(docx_bytes)} bytes")
        
        docx2pdf_success = False
        if DOCX2PDF_AVAILABLE:
            try:
                print("[PDF] Step 2: Trying docx2pdf conversion...")
                docx2pdf_convert(docx_path, pdf_path)
                if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 100:
                    docx2pdf_success = True
                    conversion_method = "docx2pdf"
                    print(f"[PDF] Step 2 SUCCESS: docx2pdf created PDF, size: {os.path.getsize(pdf_path)} bytes")
                else:
                    print("[PDF] Step 2 WARNING: docx2pdf did not create valid PDF")
            except Exception as e:
                print(f"[PDF] Step 2 WARNING: docx2pdf failed: {str(e)}")
                print("[PDF] Falling back to LibreOffice...")
        
        if not docx2pdf_success:
            print("[PDF] Step 2: Using LibreOffice for DOCX to PDF conversion...")
            conversion_method = "libreoffice"
            try:
                result = subprocess.run(
                    ["soffice", "--headless", "--convert-to", "pdf", "--outdir", temp_dir, docx_path],
                    capture_output=True,
                    timeout=180,
                    env={**os.environ, "HOME": temp_dir}
                )
                print(f"[PDF] Step 2: LibreOffice returned code {result.returncode}")
                if result.stderr:
                    print(f"[PDF] LibreOffice stderr: {result.stderr.decode()[:500]}")
            except subprocess.TimeoutExpired:
                print("[PDF] ERROR: LibreOffice timed out after 180s")
                raise Exception("DOCX to PDF conversion timed out (180s). Document may be too large.")
            except FileNotFoundError:
                print("[PDF] ERROR: LibreOffice/soffice not found")
                raise Exception("LibreOffice is not installed. Cannot convert DOCX to PDF.")
        
        if not os.path.exists(pdf_path):
            files_in_dir = os.listdir(temp_dir)
            print(f"[PDF] ERROR: PDF not created. Files in temp dir: {files_in_dir}")
            raise Exception(f"PDF was not created by LibreOffice. Return code: {result.returncode}")
        
        pdf_size = os.path.getsize(pdf_path)
        print(f"[PDF] Step 2 complete: PDF created, size: {pdf_size} bytes")
        
        if pdf_size < 100:
            raise Exception("Generated PDF is too small, likely corrupted.")
        
        try:
            print(f"[PDF] Step 3: Converting PDF to images at {dpi} DPI...")
            images = convert_from_path(
                pdf_path,
                dpi=dpi,
                fmt="jpeg",
                thread_count=1,
                use_pdftocairo=True
            )
            print(f"[PDF] Step 3 complete: Converted {len(images)} pages to images")
        except Exception as e:
            print(f"[PDF] ERROR in Step 3: {str(e)}")
            raise Exception(f"Failed to convert PDF to images: {str(e)}")
        
        if not images or len(images) == 0:
            raise Exception("No pages were extracted from the PDF.")
        
        image_paths = []
        for i, img in enumerate(images):
            img_path = os.path.join(temp_dir, f"page_{i:03d}.jpg")
            img = img.convert("RGB")
            img.save(img_path, "JPEG", quality=95, optimize=True)
            image_paths.append(img_path)
        print(f"[PDF] Step 4: Saved {len(image_paths)} JPEG images")
        
        try:
            print("[PDF] Step 5: Merging images into final PDF...")
            final_pdf_bytes = img2pdf.convert(image_paths)
            print(f"[PDF] Step 5 complete: Final PDF size: {len(final_pdf_bytes)} bytes")
        except Exception as e:
            print(f"[PDF] ERROR in Step 5: {str(e)}")
            raise Exception(f"Failed to merge images into PDF: {str(e)}")
        
        if not final_pdf_bytes or len(final_pdf_bytes) < 500:
            raise Exception("Final PDF is empty or too small.")
        
        print(f"[PDF] SUCCESS: Generated {len(final_pdf_bytes)} byte image-based PDF using {conversion_method}")
        return final_pdf_bytes, conversion_method


class ConvertToPdfRequest(BaseModel):
    docx_base64: str
    file_name: Optional[str] = "document"


@app.post("/convert-to-pdf")
async def convert_to_pdf(body: ConvertToPdfRequest):
    """Convert a base64-encoded DOCX to an image-based PDF (non-editable)"""
    try:
        docx_bytes = base64.b64decode(body.docx_base64)
        
        pdf_bytes, conversion_method = convert_docx_to_image_pdf(docx_bytes)
        
        pdf_base64 = base64.b64encode(pdf_bytes).decode("utf-8")
        
        return {
            "success": True,
            "pdf_base64": pdf_base64,
            "file_name": f"{body.file_name}.pdf",
            "file_size": len(pdf_bytes),
            "conversion_method": conversion_method
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"PDF conversion failed: {str(e)}")


class DownloadDocumentRequest(BaseModel):
    template_name: str
    vessel_imo: Optional[str] = None
    vessel_id: Optional[int] = None
    buyer_id: Optional[str] = None
    seller_id: Optional[str] = None
    product_id: Optional[str] = None
    refinery_id: Optional[str] = None
    departure_port_id: Optional[int] = None
    destination_port_id: Optional[int] = None
    buyer_bank_id: Optional[str] = None
    seller_bank_id: Optional[str] = None


@app.post("/download-document")
async def download_document(body: DownloadDocumentRequest):
    """Generate and download document directly as PDF or DOCX file (not base64 JSON)"""
    if not supabase:
        raise HTTPException(status_code=500, detail="Database not configured")
    
    template_name = body.template_name
    if not template_name.lower().endswith('.docx'):
        template_name = f"{template_name}.docx"
    
    file_path = os.path.join(TEMPLATES_DIR, template_name)
    if not os.path.exists(file_path):
        file_path_no_ext = os.path.join(TEMPLATES_DIR, body.template_name)
        if os.path.exists(file_path_no_ext):
            file_path = file_path_no_ext
            template_name = body.template_name
        else:
            raise HTTPException(status_code=404, detail=f"Template '{body.template_name}' not found")
    
    vessel_id = body.vessel_id
    if body.vessel_imo and not vessel_id:
        try:
            response = supabase.table("vessels").select("id").eq("imo", body.vessel_imo).execute()
            if response.data and len(response.data) > 0:
                vessel_id = response.data[0]["id"]
        except Exception:
            pass
    
    doc = Document(file_path)
    placeholders_by_location = extract_placeholders_from_document(doc)
    all_placeholders = set()
    for phs in placeholders_by_location.values():
        all_placeholders.update(phs)
    
    placeholder_data = {}
    context = {}
    
    if vessel_id:
        try:
            response = supabase.table("vessels").select("*").eq("id", vessel_id).execute()
            if response.data:
                vessel = response.data[0]
                context["vessel_name"] = vessel.get("name", "")
                for key, value in vessel.items():
                    placeholder_data[f"vessel_{key}"] = value
                    placeholder_data[key] = value
        except Exception:
            pass
    
    if body.buyer_id:
        try:
            response = supabase.table("buyer_companies").select("*").eq("id", body.buyer_id).execute()
            if response.data:
                for key, value in response.data[0].items():
                    placeholder_data[f"buyer_{key}"] = value
        except Exception:
            pass
    
    if body.seller_id:
        try:
            response = supabase.table("seller_companies").select("*").eq("id", body.seller_id).execute()
            if response.data:
                for key, value in response.data[0].items():
                    placeholder_data[f"seller_{key}"] = value
        except Exception:
            pass
    
    missing_placeholders = list(all_placeholders - set(placeholder_data.keys()))
    if missing_placeholders and openai_client:
        ai_values = await generate_ai_values(missing_placeholders, context)
        for ph, value in ai_values.items():
            placeholder_data[ph] = value
    
    doc = Document(file_path)
    replace_placeholders_in_document(doc, placeholder_data)
    
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    docx_bytes = output.read()
    
    base_name = template_name.replace(".docx", "").replace(".DOCX", "")
    
    pdf_bytes, conversion_method = convert_docx_to_image_pdf(docx_bytes, dpi=200)
    print(f"[DOWNLOAD] PDF generated using {conversion_method}")
    return Response(
        content=pdf_bytes,
        media_type="application/pdf",
        headers={
            "Content-Disposition": f'attachment; filename="{base_name}_filled.pdf"',
            "Access-Control-Expose-Headers": "Content-Disposition"
        }
    )


@app.post("/process-document-pdf")
async def process_document_and_convert_to_pdf(body: ProcessDocumentRequest):
    """Process document and return as image-based PDF"""
    if not supabase:
        raise HTTPException(status_code=500, detail="Database not configured")
    
    file_path = os.path.join(TEMPLATES_DIR, body.template_name)
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="Template not found")
    
    doc = Document(file_path)
    placeholders_by_location = extract_placeholders_from_document(doc)
    all_placeholders = set()
    for phs in placeholders_by_location.values():
        all_placeholders.update(phs)
    
    placeholder_data = {}
    from_database = 0
    from_ai = 0
    context = {}
    
    if body.vessel_id:
        try:
            response = supabase.table("vessels").select("*").eq("id", body.vessel_id).execute()
            if response.data:
                vessel = response.data[0]
                context["vessel_name"] = vessel.get("name", "")
                for key, value in vessel.items():
                    placeholder_data[f"vessel_{key}"] = value
                    placeholder_data[key] = value
                from_database += len(vessel)
        except Exception as e:
            print(f"Error fetching vessel: {e}")
    
    if body.buyer_id:
        try:
            response = supabase.table("buyer_companies").select("*").eq("id", body.buyer_id).execute()
            if response.data:
                buyer = response.data[0]
                context["buyer_name"] = buyer.get("name", "")
                for key, value in buyer.items():
                    placeholder_data[f"buyer_{key}"] = value
                from_database += len(buyer)
        except Exception as e:
            print(f"Error fetching buyer: {e}")
    
    if body.seller_id:
        try:
            response = supabase.table("seller_companies").select("*").eq("id", body.seller_id).execute()
            if response.data:
                seller = response.data[0]
                context["seller_name"] = seller.get("name", "")
                for key, value in seller.items():
                    placeholder_data[f"seller_{key}"] = value
                from_database += len(seller)
        except Exception as e:
            print(f"Error fetching seller: {e}")
    
    if body.product_id:
        try:
            response = supabase.table("oil_products").select("*").eq("id", body.product_id).execute()
            if response.data:
                product = response.data[0]
                context["cargo_type"] = product.get("commodity_name", "")
                for key, value in product.items():
                    placeholder_data[f"product_{key}"] = value
                from_database += len(product)
        except Exception as e:
            print(f"Error fetching product: {e}")
    
    if body.refinery_id:
        try:
            response = supabase.table("refineries").select("*").eq("id", body.refinery_id).execute()
            if response.data:
                refinery = response.data[0]
                for key, value in refinery.items():
                    placeholder_data[f"refinery_{key}"] = value
                from_database += len(refinery)
        except Exception as e:
            print(f"Error fetching refinery: {e}")
    
    if body.departure_port_id:
        try:
            response = supabase.table("ports").select("*").eq("id", body.departure_port_id).execute()
            if response.data:
                port = response.data[0]
                context["departure_port"] = port.get("name", "")
                for key, value in port.items():
                    placeholder_data[f"departure_port_{key}"] = value
                from_database += len(port)
        except Exception as e:
            print(f"Error fetching departure port: {e}")
    
    if body.destination_port_id:
        try:
            response = supabase.table("ports").select("*").eq("id", body.destination_port_id).execute()
            if response.data:
                port = response.data[0]
                context["destination_port"] = port.get("name", "")
                for key, value in port.items():
                    placeholder_data[f"destination_port_{key}"] = value
                from_database += len(port)
        except Exception as e:
            print(f"Error fetching destination port: {e}")
    
    if body.buyer_bank_id:
        try:
            response = supabase.table("buyer_company_bank_accounts").select("*").eq("id", body.buyer_bank_id).execute()
            if response.data:
                bank = response.data[0]
                for key, value in bank.items():
                    placeholder_data[f"buyer_bank_{key}"] = value
                from_database += len(bank)
        except Exception as e:
            print(f"Error fetching buyer bank: {e}")
    
    if body.seller_bank_id:
        try:
            response = supabase.table("seller_company_bank_accounts").select("*").eq("id", body.seller_bank_id).execute()
            if response.data:
                bank = response.data[0]
                for key, value in bank.items():
                    placeholder_data[f"seller_bank_{key}"] = value
                from_database += len(bank)
        except Exception as e:
            print(f"Error fetching seller bank: {e}")
    
    missing_placeholders = []
    for ph in all_placeholders:
        found = False
        for key in placeholder_data:
            if ph.lower() == key.lower() or ph.lower().replace("_", "") == key.lower().replace("_", ""):
                found = True
                break
        if not found:
            missing_placeholders.append(ph)
    
    if missing_placeholders and openai_client:
        ai_values = await generate_ai_values(missing_placeholders, context)
        for ph, value in ai_values.items():
            placeholder_data[ph] = value
            from_ai += 1
        missing_placeholders = [p for p in missing_placeholders if p not in ai_values]
    
    doc = Document(file_path)
    replacements_made = replace_placeholders_in_document(doc, placeholder_data)
    
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    docx_bytes = output.read()
    
    pdf_bytes, conversion_method = convert_docx_to_image_pdf(docx_bytes, dpi=200)
    pdf_base64 = base64.b64encode(pdf_bytes).decode("utf-8")
    
    output_filename = body.template_name.replace(".docx", "_filled.pdf")
    
    return {
        "success": True,
        "pdf_base64": pdf_base64,
        "file_name": output_filename,
        "replacements_made": replacements_made,
        "from_database": from_database,
        "from_ai": from_ai,
        "conversion_method": conversion_method,
        "missing_placeholders": missing_placeholders
    }


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=5000)
